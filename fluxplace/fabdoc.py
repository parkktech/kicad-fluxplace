#!/usr/bin/env python3
"""Render a FAB-SUBMISSION markdown brief as a styled Word document.

The person who uploads to PCBWay/JLCPCB is not the person who drew the board:
they want something readable outside a zip, on a phone if need be. This keeps
the .md as the single source of truth and generates the .docx from it, so the
two can never disagree (the V1.2 pair drifted by hand).

Requires python-docx; `deliver` degrades to markdown-only without it. The import
is deliberately INSIDE build(): the CLI often runs on KiCad's python (for
pcbnew), which usually has no python-docx, and `render()` below needs to be
importable there so it can hand the job to a python that does.

Used by `fluxplace deliver` and `fluxplace/pcbway.py`; also runnable directly:
    python3 -m fluxplace.fabdoc IN.md OUT.docx "Title" "Subtitle"
"""
import os
import re
import subprocess
import sys

# plain tuples, not RGBColor: this module must import without python-docx so
# render() can still reach an interpreter that has it
NAVY = (0x1F, 0x3A, 0x5F)
GREY = (0x60, 0x60, 0x60)
RED = (0xB0, 0x20, 0x20)


def _runs(par, text, size=10.5, bold=False, color=None):
    """Emit text, honouring **bold** and `code` spans."""
    from docx.shared import Pt, RGBColor
    for chunk in re.split(r"(\*\*[^*]+\*\*|\*[^*]+\*|`[^`]+`)", text):
        if not chunk:
            continue
        if chunk.startswith("*") and not chunk.startswith("**") and \
                chunk.endswith("*"):
            r = par.add_run(chunk[1:-1]); r.italic = True; r.bold = bold
        elif chunk.startswith("**") and chunk.endswith("**"):
            # a bold span may contain `code`; Word has no nesting here, so
            # keep the bold and drop the backticks rather than print them
            r = par.add_run(chunk[2:-2].replace("`", "")); r.bold = True
        elif chunk.startswith("`") and chunk.endswith("`"):
            r = par.add_run(chunk[1:-1]); r.font.name = "Consolas"
            r.font.size = Pt(size - 1)
        else:
            r = par.add_run(chunk); r.bold = bold
        r.font.size = Pt(size)
        if color is not None:
            r.font.color.rgb = RGBColor(*color)
    return par


def build(md_path, out_path, title, subtitle):
    import docx
    from docx.shared import Inches, Pt
    lines = open(md_path, encoding="utf-8").read().splitlines()
    d = docx.Document()
    s = d.sections[0]
    s.left_margin = s.right_margin = Inches(0.8)
    s.top_margin = s.bottom_margin = Inches(0.7)

    p = d.add_paragraph(); _runs(p, title, size=22, bold=True, color=NAVY)
    p.paragraph_format.space_after = Pt(2)
    p = d.add_paragraph(); _runs(p, subtitle, size=13, color=GREY)
    p.paragraph_format.space_after = Pt(10)

    i = 0
    # the markdown's own H1 + lead-in lines become the grey meta block
    meta = []
    while i < len(lines) and not lines[i].startswith("## "):
        ln = lines[i].strip()
        if ln.startswith("# "):
            i += 1
            continue
        if ln:
            meta.append(ln)
        elif meta:
            p = d.add_paragraph()
            _runs(p, re.sub(r"\*\*", "", " ".join(meta)), size=9.5, color=GREY)
            p.paragraph_format.space_after = Pt(3)
            meta = []
        i += 1
    if meta:
        p = d.add_paragraph()
        _runs(p, re.sub(r"\*\*", "", " ".join(meta)), size=9.5, color=GREY)
        p.paragraph_format.space_after = Pt(3)

    while i < len(lines):
        ln = lines[i]
        st = ln.strip()

        if st.startswith("## "):
            p = d.add_paragraph()
            _runs(p, st[3:].replace("§", "").strip(), size=14, bold=True,
                  color=NAVY)
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(4)
            i += 1
            continue

        # markdown table -> Table Grid
        if st.startswith("|") and i + 1 < len(lines) and \
                re.match(r"^\s*\|[\s:|-]+\|\s*$", lines[i + 1]):
            rows = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                cells = [c.strip() for c in lines[i].strip().strip("|").split("|")]
                if not re.match(r"^[\s:|-]+$", "".join(cells)):
                    rows.append(cells)
                i += 1
            if rows:
                cols = max(len(r) for r in rows)
                t = d.add_table(rows=0, cols=cols)
                t.style = "Table Grid"
                for ri, row in enumerate(rows):
                    cs = t.add_row().cells
                    for ci in range(cols):
                        txt = row[ci] if ci < len(row) else ""
                        cell_p = cs[ci].paragraphs[0]
                        _runs(cell_p, txt, size=9,
                              bold=(ri == 0), color=(NAVY if ri == 0 else None))
                d.add_paragraph().paragraph_format.space_after = Pt(6)
            continue

        # blockquote -> indented boxed note (the paste-to-engineer text)
        if st.startswith(">"):
            buf = []
            while i < len(lines) and lines[i].strip().startswith(">"):
                buf.append(lines[i].strip().lstrip(">").strip())
                i += 1
            text = " ".join(x for x in buf if x)
            t = d.add_table(rows=1, cols=1)
            t.style = "Table Grid"
            _runs(t.rows[0].cells[0].paragraphs[0], text, size=9.5)
            d.add_paragraph().paragraph_format.space_after = Pt(6)
            continue

        if st.startswith("- ") or st.startswith("* "):
            buf = [st[2:]]
            i += 1
            while i < len(lines):            # absorb wrapped continuation
                cur = lines[i].strip()
                if not cur or cur.startswith(("## ", "|", ">", "- ", "* ")):
                    break
                buf.append(cur)
                i += 1
            p = d.add_paragraph(style="List Bullet")
            _runs(p, " ".join(buf), size=10)
            p.paragraph_format.space_after = Pt(2)
            continue

        if st:
            # markdown hard-wraps prose; join a run of plain lines into ONE
            # Word paragraph or the output reads as ragged half-sentences
            buf = []
            while i < len(lines):
                cur = lines[i].strip()
                if (not cur or cur.startswith(("## ", "|", ">", "- ", "* "))):
                    break
                buf.append(cur)
                i += 1
            text = " ".join(buf)
            p = d.add_paragraph()
            colour = RED if text.startswith("**Do not") or "BLOCKING" in text else None
            _runs(p, text, size=10.5, color=colour)
            p.paragraph_format.space_after = Pt(5)
            continue
        i += 1

    d.save(out_path)
    return out_path


def render(md_path, out_path, title, subtitle):
    """build(), but never raises: -> the .docx path, or None if no interpreter
    on this machine has python-docx.

    The CLI runs under KiCad's python for pcbnew, which usually lacks
    python-docx — so rather than silently dropping the Word file (the one file
    the buyer actually opens), hand the job to a python that has it.
    """
    try:
        build(md_path, out_path, title, subtitle)
        return out_path
    except ImportError:
        pass
    here = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
    for py in ("python3", "/usr/bin/python3"):
        try:
            r = subprocess.run([py, "-m", "fluxplace.fabdoc", md_path,
                                out_path, title, subtitle], cwd=here,
                               capture_output=True, timeout=120)
            if r.returncode == 0 and os.path.exists(out_path):
                return out_path
        except Exception:
            continue
    return None


if __name__ == "__main__":
    md, out, title, subtitle = sys.argv[1:5]
    print("wrote", build(md, out, title, subtitle))
